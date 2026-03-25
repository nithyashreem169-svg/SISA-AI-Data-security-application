"""File Parser - Parses different file types into lines"""
from io import BytesIO
from app.utils.logger import logger

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class FileParser:
    """Parse uploaded files into structured format"""
    
    @staticmethod
    def parse_txt_log(file_content: bytes) -> tuple[bool, list[str], str]:
        """
        Parse TXT/LOG files line by line
        
        Args:
            file_content: Raw file bytes
            
        Returns:
            (success, lines_list, error_message)
        """
        try:
            # Decode bytes to string
            text = file_content.decode('utf-8', errors='ignore')
            
            # Split into lines
            lines = text.strip().split('\n')
            
            logger.info(f"Parsed file: {len(lines)} lines")
            return True, lines, ""
            
        except Exception as e:
            error_msg = f"Error parsing TXT/LOG file: {str(e)}"
            logger.error(error_msg)
            return False, [], error_msg
    
    @staticmethod
    def parse_docx(file_content: bytes) -> tuple[bool, list[str], str]:
        """
        Parse DOCX files (.docx - Office Open XML format)
        
        Args:
            file_content: Raw file bytes
            
        Returns:
            (success, lines_list, error_message)
        """
        try:
            if not DOCX_AVAILABLE:
                return False, [], "python-docx not installed. Install with: pip install python-docx"
            
            # Read DOCX from bytes
            docx_buffer = BytesIO(file_content)
            doc = Document(docx_buffer)
            
            # Extract all paragraphs and tables
            lines = []
            
            # Get paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Only add non-empty lines
                    lines.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        lines.append(row_text)
            
            logger.info(f"Parsed DOCX file: {len(lines)} lines/rows extracted")
            return True, lines, ""
            
        except Exception as e:
            error_msg = f"Error parsing DOCX file: {str(e)}"
            logger.error(error_msg)
            return False, [], error_msg
    
    @staticmethod
    def parse_doc(file_content: bytes) -> tuple[bool, list[str], str]:
        """
        Parse DOC files (.doc - older format, use python-docx)
        
        Args:
            file_content: Raw file bytes
            
        Returns:
            (success, lines_list, error_message)
        """
        try:
            if not DOCX_AVAILABLE:
                return False, [], "python-docx not installed. Install with: pip install python-docx"
            
            # python-docx can also handle some .doc files (depends on format)
            # For true .doc format, consider using python-docx2txt or similar
            docx_buffer = BytesIO(file_content)
            
            try:
                doc = Document(docx_buffer)
                lines = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        lines.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        if row_text.strip():
                            lines.append(row_text)
                
                logger.info(f"Parsed DOC file: {len(lines)} lines/rows extracted")
                return True, lines, ""
                
            except Exception as inner_e:
                # If python-docx fails, try basic text extraction (fallback)
                logger.warning(f"python-docx parsing failed for DOC: {inner_e}. Attempting fallback...")
                
                # Basic fallback: extract readable text from bytes
                text = file_content.decode('utf-8', errors='ignore')
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                
                if lines:
                    logger.info(f"DOC parsed via fallback: {len(lines)} lines extracted")
                    return True, lines, ""
                else:
                    return False, [], "Could not extract text from DOC file"
                    
        except Exception as e:
            error_msg = f"Error parsing DOC file: {str(e)}"
            logger.error(error_msg)
            return False, [], error_msg
    
    @staticmethod
    def parse_file(file_content: bytes, filename: str) -> tuple[bool, list[str], str]:
        """
        Main file parser dispatcher
        
        Args:
            file_content: Raw file bytes
            filename: Name of file (determines format)
            
        Returns:
            (success, lines_list, error_message)
        """
        ext = filename.split('.')[-1].lower()
        
        if ext in ['txt', 'log']:
            return FileParser.parse_txt_log(file_content)
        elif ext == 'docx':
            return FileParser.parse_docx(file_content)
        elif ext == 'doc':
            return FileParser.parse_doc(file_content)
        else:
            error_msg = f"Parser not yet implemented for .{ext}"
            logger.warning(error_msg)
            return False, [], error_msg
