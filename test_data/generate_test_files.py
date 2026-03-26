"""
Generate test files for SISA AI Security Platform testing.
Creates: .txt, .log, .docx, and .doc files with a mix of safe and unsafe content
that triggers all 15 detection patterns.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# FILE 1: security_audit_report.txt
# A fake security audit report with embedded sensitive data
# ============================================================
txt_content = """=== ACME Corp Internal Security Audit Report ===
Date: 2026-03-25
Auditor: security-team@acmecorp.com
Classification: CONFIDENTIAL

--- Section 1: System Overview ---
Server: prod-web-01 (192.168.1.100)
Server: prod-db-01 (10.0.0.55)
Server: staging-api (172.16.0.12)
Application: ACME Customer Portal v4.2.1
Environment: Production

--- Section 2: Access Control Findings ---
[SAFE] All admin accounts use MFA.
[SAFE] Role-based access control is properly configured.
[SAFE] Session timeout is set to 15 minutes.

[FINDING] Hardcoded admin credentials found in deployment script:
  username=admin_deploy
  password=Pr0d_D3pl0y!S3cret#2026

[FINDING] Database connection string exposed in application.properties:
  postgresql://app_user:DbP@ss2026!@prod-db-01:5432/acme_customers

[FINDING] Legacy config contains AWS credentials:
  FAKE_AKIA4EXAMPLE1234ABCD

--- Section 3: Data Exposure Audit ---
[SAFE] Customer data is encrypted at rest with AES-256.
[SAFE] TLS 1.3 enforced on all endpoints.

[FINDING] Customer PII found in debug log output:
  Customer SSN: 451-78-9023
  Contact: john.doe@example.com, Phone: (555) 867-5309
  Backup contact: jane.smith@acmecorp.com

[FINDING] Payment processing log contains card data:
  Transaction #TXN-9281: Card 4532-0151-2832-9744 processed successfully
  Transaction #TXN-9282: Card 5425 2334 6789 0123 declined - insufficient funds

--- Section 4: Application Security ---
[SAFE] Input validation is applied on all form fields.
[SAFE] CSRF tokens are properly implemented.

[FINDING] API key exposed in frontend JavaScript bundle:
  api_key=fake_sk_live_4eC39HqLyjWDarjtT1zdp7dc_xJ29mK4pLqR8nYz

[FINDING] GitHub personal access token found in CI/CD config:
  fake_ghp_aBcDeFgHiJkLmNoPqRsTuVwXyZ1234567890

[FINDING] Error response leaking internal paths:
  Traceback (most recent call last):
    File "/home/deploy/acme-portal/app/services/payment.py", line 142
  exception: NullPointerException - payment gateway returned null for customer_id=98231

[FINDING] Debug mode enabled in production:
  debug=true in /etc/acme/portal.conf

--- Section 5: Recommendations ---
1. Rotate all exposed credentials immediately
2. Implement secrets management (HashiCorp Vault)
3. Add log sanitization pipeline
4. Disable debug mode in production
5. Remove hardcoded credentials from codebase

=== END OF REPORT ===
"""

# ============================================================
# FILE 2: application_server.log
# A realistic server log with security events
# ============================================================
log_content = """2026-03-25 08:00:01 INFO  [main] Application starting on port 8443...
2026-03-25 08:00:02 INFO  [main] Loading configuration from /etc/acme/app.conf
2026-03-25 08:00:03 INFO  [main] Database connection established successfully
2026-03-25 08:00:03 INFO  [main] Server ready, accepting connections
2026-03-25 08:15:22 INFO  [auth] User login successful: user_id=1024
2026-03-25 08:15:45 INFO  [auth] User login successful: user_id=2048
2026-03-25 08:32:10 WARNING [auth] Failed login attempt for user: admin@acmecorp.com
2026-03-25 08:32:15 WARNING [auth] Failed login attempt for user: admin@acmecorp.com
2026-03-25 08:32:19 WARNING [auth] Failed login attempt for user: admin@acmecorp.com
2026-03-25 08:32:22 WARNING [auth] Failed login attempt for user: admin@acmecorp.com
2026-03-25 08:32:25 WARNING [auth] Failed login attempt for user: admin@acmecorp.com
2026-03-25 08:32:28 WARNING [auth] Failed login attempt for user: admin@acmecorp.com
2026-03-25 08:32:30 ERROR  [auth] Account locked due to repeated failed attempts: admin@acmecorp.com
2026-03-25 09:01:44 ERROR  [payment] Payment processing failed for customer_id=55123
2026-03-25 09:01:44 ERROR  [payment] Traceback (most recent call last):
2026-03-25 09:01:44 ERROR  [payment]   File "/home/deploy/acme-portal/services/payment_handler.py", line 89
2026-03-25 09:01:44 ERROR  [payment]   at com.acme.payment.Gateway.processTransaction(Gateway.java:241)
2026-03-25 09:01:44 ERROR  [payment] exception: ConnectionRefusedException - payment gateway timeout after 30s
2026-03-25 09:12:33 WARNING [api] Slow query detected: SELECT * FROM customers WHERE ssn='329-44-8827' (took 4.2s)
2026-03-25 09:12:33 WARNING [api] Query exposed SSN: 329-44-8827
2026-03-25 09:30:00 INFO  [scheduler] Running daily backup job
2026-03-25 09:30:01 INFO  [scheduler] Connecting to backup server at 10.0.0.200
2026-03-25 09:45:12 ERROR  [config] Configuration reload failed - using cached config
2026-03-25 09:45:12 ERROR  [config] Reason: password=B@ckupS3rv3r!2026 rejected by vault
2026-03-25 10:00:05 INFO  [api] Health check passed - all services operational
2026-03-25 10:15:33 WARNING [security] Suspicious request from IP 203.0.113.42: attempted path traversal
2026-03-25 10:15:33 WARNING [security] Request contained: GET /../../etc/passwd HTTP/1.1
2026-03-25 10:30:22 ERROR  [api] Unhandled exception in /api/v2/users endpoint
2026-03-25 10:30:22 ERROR  [api] Traceback (most recent call last):
2026-03-25 10:30:22 ERROR  [api]   File "/home/deploy/acme-portal/api/users.py", line 156
2026-03-25 10:30:22 ERROR  [api] fatal error: SIGSEGV received in worker process pid=44821
2026-03-25 11:00:00 INFO  [audit] Daily audit: 1,247 requests processed, 3 blocked
2026-03-25 11:05:18 WARNING [deploy] Deployment script using hardcoded credentials:
2026-03-25 11:05:18 WARNING [deploy] username=svc_deploy
2026-03-25 11:05:18 WARNING [deploy] password=D3pl0ym3nt_K3y!X9z
2026-03-25 11:05:18 WARNING [deploy] secret_key=fake_rk_live_9a8b7c6d5e4f3g2h1i0jklmnopqrstuv
2026-03-25 11:20:00 INFO  [api] API key rotation reminder: current key api_key=ak_prod_Zy9XwV8uT7sR6qP5oN4mL3k
2026-03-25 11:45:33 ERROR  [db] Connection pool exhausted - 50/50 connections in use
2026-03-25 11:45:33 ERROR  [db] mongodb://analytics_svc:M0ng0P@ss!@analytics-db:27017/metrics
2026-03-25 12:00:00 INFO  [scheduler] Sending daily report to ops-team@acmecorp.com
2026-03-25 12:00:01 INFO  [scheduler] Report includes customer contact: sarah.johnson@clientco.com, (415) 555-0198
2026-03-25 12:30:15 INFO  [main] Graceful shutdown initiated
2026-03-25 12:30:16 INFO  [main] All connections closed. Server stopped.
"""

# ============================================================
# FILE 3: employee_data_export.docx
# A Word document simulating an employee data export
# ============================================================
def create_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('ACME Corp - Employee Data Export', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Classification: INTERNAL - RESTRICTED')
    doc.add_paragraph('Generated: 2026-03-25 | Export ID: EXP-2026-0325-001')
    doc.add_paragraph('')
    
    # Section 1: Safe content
    doc.add_heading('1. Department Summary', level=1)
    doc.add_paragraph('Total Employees: 342')
    doc.add_paragraph('Departments: Engineering (89), Sales (67), Marketing (45), HR (32), Finance (28), Operations (81)')
    doc.add_paragraph('Average tenure: 3.2 years')
    doc.add_paragraph('Remote workers: 45%')
    
    # Section 2: Employee records with PII
    doc.add_heading('2. Employee Records', level=1)
    doc.add_paragraph('WARNING: This section contains personally identifiable information (PII).')
    doc.add_paragraph('')
    
    # Employee 1
    doc.add_heading('Employee #1001 - Robert Chen', level=2)
    doc.add_paragraph('Department: Engineering')
    doc.add_paragraph('Email: robert.chen@acmecorp.com')
    doc.add_paragraph('Phone: (408) 555-0147')
    doc.add_paragraph('SSN: 512-63-8847')
    doc.add_paragraph('Emergency Contact: lisa.chen@gmail.com, (408) 555-0291')
    
    # Employee 2
    doc.add_heading('Employee #1002 - Maria Santos', level=2)
    doc.add_paragraph('Department: Finance')
    doc.add_paragraph('Email: maria.santos@acmecorp.com')
    doc.add_paragraph('Phone: (650) 555-0183')
    doc.add_paragraph('SSN: 287-41-5529')
    doc.add_paragraph('Direct Deposit: Card ending 4716-8234-1092-5567')
    
    # Employee 3
    doc.add_heading('Employee #1003 - James Wilson', level=2)
    doc.add_paragraph('Department: IT Operations')
    doc.add_paragraph('Email: james.wilson@acmecorp.com')
    doc.add_paragraph('Phone: (212) 555-0165')
    doc.add_paragraph('SSN: 634-92-1178')
    
    # Section 3: IT credentials (BAD PRACTICE)
    doc.add_heading('3. System Access Credentials', level=1)
    doc.add_paragraph('NOTE: These credentials should NOT be in this document!')
    doc.add_paragraph('')
    doc.add_paragraph('VPN Access:')
    doc.add_paragraph('  username=vpn_admin')
    doc.add_paragraph('  password=VPN@ccess2026!Secure')
    doc.add_paragraph('')
    doc.add_paragraph('AWS Production Console:')
    doc.add_paragraph('  Access Key: FAKE_AKIA3EXAMPLE5678WXYZ')
    doc.add_paragraph('  api_secret=AwsS3cr3tK3y_Pr0duct10n_2026_xYzAbCdEf')
    doc.add_paragraph('')
    doc.add_paragraph('Database Admin:')
    doc.add_paragraph('  mysql://db_admin:Sql@dm1n!2026@prod-mysql.acme.internal:3306/employees')
    
    # Section 4: Financial data
    doc.add_heading('4. Corporate Card Information', level=1)
    doc.add_paragraph('Department Head corporate cards for expense reporting:')
    doc.add_paragraph('')
    
    # Table with card data
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Name', 'Department', 'Card Number', 'Exp Date']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    data = [
        ['R. Chen', 'Engineering', '4532 7891 2345 6780', '08/27'],
        ['M. Santos', 'Finance', '5425 1234 5678 9012', '11/28'],
        ['J. Wilson', 'IT Ops', '4716 9876 5432 1098', '03/27'],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value
    
    # Section 5: Safe summary
    doc.add_paragraph('')
    doc.add_heading('5. Export Notes', level=1)
    doc.add_paragraph('This export was generated for the annual HR compliance review.')
    doc.add_paragraph('All data must be handled according to ACME Corp Data Handling Policy v3.1.')
    doc.add_paragraph('Retain for 90 days, then securely destroy.')
    doc.add_paragraph('Questions: Contact HR Security at hrsecurity@acmecorp.com')
    
    filepath = os.path.join(OUTPUT_DIR, 'employee_data_export.docx')
    doc.save(filepath)
    print(f"Created: {filepath}")


# ============================================================
# FILE 4: incident_response_notes.doc (saved as docx internally)
# A document about a security incident
# ============================================================
def create_doc():
    doc = Document()
    
    title = doc.add_heading('Security Incident Response - INC-2026-0042', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Severity: P1 - CRITICAL')
    doc.add_paragraph('Status: Under Investigation')
    doc.add_paragraph('Reported: 2026-03-24 14:32 UTC')
    doc.add_paragraph('')
    
    doc.add_heading('Incident Summary', level=1)
    doc.add_paragraph(
        'On March 24, 2026, our monitoring system detected unauthorized API calls '
        'originating from IP address 198.51.100.78. The attacker appeared to use '
        'compromised credentials to access the customer database.'
    )
    
    doc.add_heading('Compromised Credentials', level=1)
    doc.add_paragraph('The following credentials were confirmed compromised:')
    doc.add_paragraph('')
    doc.add_paragraph('Service Account:')
    doc.add_paragraph('  username=svc_customer_api')
    doc.add_paragraph('  password=Cust0m3r_AP1_2026!xYz')
    doc.add_paragraph('')
    doc.add_paragraph('API Token found in attacker request logs:')
    doc.add_paragraph('  auth_token=fake_pat-ghk83mf92nxk47vb29sl05pa61wq8rti')
    doc.add_paragraph('')
    doc.add_paragraph('The attacker also obtained an internal Slack webhook:')
    doc.add_paragraph('  access_token=fake_xoxb-12345678901-abcdefghijklm_nopqrstuvwxyz')
    
    doc.add_heading('Affected Data', level=1)
    doc.add_paragraph('Customer records potentially exfiltrated:')
    doc.add_paragraph('')
    doc.add_paragraph('Customer: Alice Thompson')
    doc.add_paragraph('  Email: alice.thompson@example.com')
    doc.add_paragraph('  Phone: (312) 555-0234')
    doc.add_paragraph('  SSN: 178-52-9934')
    doc.add_paragraph('  Payment Card: 4111 1111 1111 1111')
    doc.add_paragraph('')
    doc.add_paragraph('Customer: David Park')
    doc.add_paragraph('  Email: david.park@testmail.com')
    doc.add_paragraph('  Phone: (617) 555-0189')
    doc.add_paragraph('  SSN: 623-41-8856')
    doc.add_paragraph('  Payment Card: 5500-0000-0000-0004')
    
    doc.add_heading('Technical Details', level=1)
    doc.add_paragraph('Attack timeline and error logs from affected server:')
    doc.add_paragraph('')
    doc.add_paragraph('14:32:01 - Initial access via compromised API token from 198.51.100.78')
    doc.add_paragraph('14:32:05 - Traceback (most recent call last):')
    doc.add_paragraph('14:32:05 -   File "/var/www/api/controllers/customer.py", line 88')
    doc.add_paragraph('14:32:05 - exception: UnauthorizedAccessException - elevated privilege attempt detected')
    doc.add_paragraph('14:33:12 - Attacker escalated privileges using:')
    doc.add_paragraph('14:33:12 -   secret_key=Adm1n_Escal@t10n_K3y!2026')
    doc.add_paragraph('14:35:00 - Database dump initiated from 10.0.0.55')
    doc.add_paragraph('14:38:22 - Exfiltration detected to external IP 203.0.113.99')
    doc.add_paragraph('14:40:00 - Incident detected by SIEM, automatic lockdown triggered')
    
    doc.add_heading('Remediation Steps', level=1)
    doc.add_paragraph('1. All compromised credentials have been rotated')
    doc.add_paragraph('2. Affected IP addresses blocked at firewall level')
    doc.add_paragraph('3. Customer notification in progress (GDPR 72-hour requirement)')
    doc.add_paragraph('4. Forensic image of affected servers preserved')
    doc.add_paragraph('5. PCI-DSS breach notification filed')
    
    doc.add_heading('Contact', level=1)
    doc.add_paragraph('Incident Commander: security-lead@acmecorp.com')
    doc.add_paragraph('Legal Counsel: legal-team@acmecorp.com')
    
    # Save as .doc extension (python-docx format, which the parser handles)
    filepath = os.path.join(OUTPUT_DIR, 'incident_response_notes.doc')
    doc.save(filepath)
    print(f"Created: {filepath}")


# ============================================================
# MAIN: Generate all files
# ============================================================
if __name__ == '__main__':
    print("Generating test files for SISA AI Security Platform...\n")
    
    # TXT file
    txt_path = os.path.join(OUTPUT_DIR, 'security_audit_report.txt')
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(txt_content.strip())
    print(f"Created: {txt_path}")
    
    # LOG file
    log_path = os.path.join(OUTPUT_DIR, 'application_server.log')
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write(log_content.strip())
    print(f"Created: {log_path}")
    
    # DOCX file
    create_docx()
    
    # DOC file
    create_doc()
    
    print("\n✅ All 4 test files generated successfully!")
    print("\nFiles created:")
    print("  1. security_audit_report.txt    — Audit report with passwords, SSNs, API keys, IPs")
    print("  2. application_server.log       — Server log with credentials, errors, brute-force attempts")
    print("  3. employee_data_export.docx    — Employee data with PII, credit cards, AWS keys")
    print("  4. incident_response_notes.doc  — Incident report with compromised credentials & customer data")
    print("\nUpload any of these to the SISA frontend at http://localhost:8501")
